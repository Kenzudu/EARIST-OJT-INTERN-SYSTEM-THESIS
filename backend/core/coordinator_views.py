import re
from django.conf import settings
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from django.contrib.auth.models import User
from django.utils import timezone
from django.db.models import Q
from .models import UserRole, StudentProfile, Company, Internship, Application, DocumentTemplate, CoordinatorProfile, Message, DailyJournal, DocumentTypeConfig
from .serializers import ApplicationSerializer, CompanySerializer, UserSerializer, DailyJournalSerializer
from .permissions import role_required
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from django.core.files.base import ContentFile
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.pdfgen import canvas
import io
import os
import tempfile
try:
    import pythoncom
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

# ========================================
# HELPER FUNCTIONS
# ========================================

def fill_document_template(template_path, context):
    """
    Fill a DOCX template with context variables.
    Handles standard tags, square brackets, underlines, and Tab leaders (\t).
    """
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"FAILED TO OPEN TEMPLATE: {str(e)}")
        return None

    # DIAGNOSTIC LOGGING (Helpful for troubleshooting)
    with open('document_generation_debug.log', 'a', encoding='utf-8') as log:
        log.write(f"\n--- Processing Template: {template_path} ---\n")
        log.write(f"Context: {context}\n")

    # Permissive regex for lines: underscores, dashes, or Tab characters (with leaders)
    # Most professional templates use \t with an underline leader
    h_line = r'([\t_]{1,}|[_\u2014\u2013-]{2,})'
    
    def process_paragraph(paragraph):
        if not paragraph.text.strip():
            return
            
        original_text = paragraph.text
        text = original_text
        
        # STEP 0: REMOVE HARDCODED EXAMPLES FROM TEMPLATE FIRST
        # Templates often have example data like "Mai Shiranui" or "Technova Company"
        # We need to replace these with the actual student's data BEFORE doing other replacements
        
        # 1. HARDCODED STUDENT NAMES (including typos found in logs)
        hardcoded_students = [
            "Mai Shiranui", "Mai Shrinaui", "Mai Shiraui", 
            "Norman Mangusin", "Rafael Kenneth Saluba", "johnsteven caspe", "John Doe"
        ]
        current_student = context.get('student_name', '')
        
        # Replace any hardcoded student name with the actual selected student
        for hardcoded in hardcoded_students:
            if hardcoded in text and hardcoded != current_student:
                with open('document_generation_debug.log', 'a', encoding='utf-8') as log:
                    log.write(f"FOUND HARDCODED STUDENT: '{hardcoded}' -> Replacing with '{current_student}'\n")
                text = text.replace(hardcoded, current_student)

        # 2. HARDCODED SECTIONS
        hardcoded_sections = ["4-C BS INFOTECH", "4-A BSCS", "3-B BS INFO TECH", "2-C BSIT"]
        current_section = context.get('section', '')
        
        # Replace any hardcoded section with the actual student's section
        for hardcoded_section in hardcoded_sections:
            if hardcoded_section in text and hardcoded_section != current_section:
                if current_section and not current_section.startswith('_'):
                    text = text.replace(hardcoded_section, current_section)
        
        # 3. HARDCODED COMPANIES
        hardcoded_companies = ["Technova Company/Agency", "Technova Company", "ABC Corporation", "XYZ Corp"]
        current_company = context.get('company_name', '')
        
        # Replace any hardcoded company name
        for hardcoded_company in hardcoded_companies:
            if hardcoded_company in text:
                if current_company and not current_company.startswith('_'):
                    text = text.replace(hardcoded_company, current_company)
                else:
                    text = text.replace(hardcoded_company, "______________________________")
        
        # 4. HARDCODED ADDRESSES
        hardcoded_addresses = [
            "21 Unang Hakbang Brgy Don Manuel Quezon City",
            "123 Main Street, City",
            "Sample Address Here"
        ]
        current_address = context.get('company_address', '')
        
        # Replace any hardcoded address
        for hardcoded_addr in hardcoded_addresses:
            if hardcoded_addr in text:
                if current_address and not current_address.startswith('_'):
                    text = text.replace(hardcoded_addr, current_address)
                else:
                    text = text.replace(hardcoded_addr, "______________________________")
        
        # 5. AGGRESSIVE REGEX CLEANUP (Catch-all for "ward [Name], [Section]")
        # Matches patterns like "ward [Anything up to ,], [Anything up to ,] to undergo"
        # Only runs if we haven't already replaced the name
        if current_student not in text and re.search(r'ward\s+[A-Z][a-z]+', text):
             pattern = r'(ward\s+)([^,]+)(,\s*)([^,]+)(\s*,?\s*to\s+undergo)'
             if re.search(pattern, text):
                 replacement = r'\1' + current_student + r'\3' + current_section + r'\5'
                 text = re.sub(pattern, replacement, text)
                 with open('document_generation_debug.log', 'a', encoding='utf-8') as log:
                     log.write(f"AGGRESSIVE REGEX REPLACEMENT APPLIED\n")
        
        # NOW continue with the regular logic...
        
        # A. CONTEXTUAL REPLACEMENTS (Based on neighboring keywords)
        # 1. Student name: "ward [tab/line]" -> "ward StudentName"
        if re.search(r'ward', text, re.IGNORECASE):
            text = re.sub(r'(ward\s*)' + h_line, r'\1' + context.get('student_name', ''), text, flags=re.IGNORECASE)
            
        # 2. Section: "[tab/line] to undergo" -> "YearSection to undergo"
        # The section placeholder comes BEFORE "to undergo"
        if re.search(r'to\s+undergo', text, re.IGNORECASE):
            section_value = context.get('section', '')
            # Replace the tab/line before "to undergo" with the section
            text = re.sub(h_line + r'(\s+to\s+undergo)', section_value + r'\1', text, flags=re.IGNORECASE)
            
        # 3. Company: "Training at [underscores/tabs]" -> "Training at CompanyName" or keep underlines
        # Look for "at" followed by underscores or tabs
        if re.search(r'at\s+[_\t]', text, re.IGNORECASE):
            company_value = context.get('company_name', '')
            # Only replace if company_value is NOT underscores (i.e., student has a real company)
            if company_value and not company_value.startswith('_'):
                text = re.sub(r'(at\s+)[_\t]+', r'\1' + company_value, text, flags=re.IGNORECASE)
            # If company_value is underlines or empty, leave the underlines as-is

        # B. EXACT MATCHES FOR EARIST TEMPLATE BRACKETS
        # These are common in the user's uploaded file even if they look like lines
        template_specific = [
            (r'\[\s*FIRST NAME MI\. SURNAME\s*\]', context.get('student_name', '')),
            (r'\[\s*NAME OF COOPERATING AGENCY\s*\]', context.get('company_name', '')),
            (r'\[\s*ADDRESS OF COOPERATING AGENCY\s*\]', context.get('company_address', '')),
            (r'\[\s*SECTION\s*\]', context.get('section', '')),
            (r'\[\s*STUDENT NAME\s*\]', context.get('student_name', '')),
            (r'\[\s*STUDENT ID\s*\]', context.get('student_id', '')),
            (r'\[\s*COURSE\s*\]', context.get('course', '')),
        ]
        for pattern, repl in template_specific:
            text = re.sub(pattern, str(repl), text, flags=re.IGNORECASE)

        # C. FALLBACK FOR ANY OTHER [KEY] OR {{key}}
        for key, val in context.items():
            label = key.upper().replace('_', ' ')
            text = re.sub(r'\[\s*' + re.escape(label) + r'\s*\]', str(val or ''), text, flags=re.IGNORECASE)
            text = re.sub(r'\{\{\s*' + re.escape(key) + r'\s*\}\}', str(val or ''), text, flags=re.IGNORECASE)

        # D. CLEANUP (Remove stagnant placeholder text from bottom areas)
        cleanup = ["FIRST NAME MI. SURNAME", "(09XX) XXX-XXXX", "(09xx) xxx-xxxx", "MARCH XX, 2024"]
        for c_text in cleanup:
            if c_text in text:
                text = text.replace(c_text, "")
        
        # E. REMOVE HARDCODED STUDENT NAMES FROM TEMPLATE
        # If the template has a specific student name hardcoded (like "Mai Shiranui"),
        # we need to replace it with the actual selected student's name
        hardcoded_names = ["Mai Shiranui", "Norman Mangusin", "Rafael Kenneth Saluba", "johnsteven caspe"]
        current_student_name = context.get('student_name', '')
        for hardcoded_name in hardcoded_names:
            if hardcoded_name in text and hardcoded_name != current_student_name:
                # Replace hardcoded name with actual student name
                text = text.replace(hardcoded_name, current_student_name)
        
        # F. ENSURE COMPANY/ADDRESS SHOW UNDERLINES WHEN EMPTY
        # If company_name is empty, make sure we show underlines
        if 'Training at' in text and context.get('company_name', '') == '':
            # Replace any existing company placeholder with clean underlines
            text = re.sub(r'(Training at\s+)[^\n,]*', r'\1______________________________', text)

        # G. FINAL PARAGRAPH UPDATE
        if text != original_text:
            paragraph.text = text
            with open('document_generation_debug.log', 'a', encoding='utf-8') as log:
                log.write(f"REPLACED: '{repr(original_text)}' WITH '{repr(text)}'\n")

    # Run for all paragraphs
    for p in doc.paragraphs:
        process_paragraph(p)
            
    # Run for all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
                    
    return doc

def fill_pdf_template(pdf_path, context):
    """
    Fills a PDF form with the provided context data.
    Requires the PDF to have form fields with names matching context keys.
    """
    try:
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Add all pages to writer
        for page in reader.pages:
            writer.add_page(page)
            
        # Try to update form fields if they exist
        # We'll try to match context keys (like student_name) to PDF field names
        fields = reader.get_fields()
        if fields:
            # Map context keys to field names (case-insensitive and partial match)
            field_data = {}
            for field_name in fields.keys():
                for key, value in context.items():
                    # Match if field name contains our key or vice versa
                    if key.lower() in field_name.lower() or field_name.lower() in key.lower():
                        field_data[field_name] = str(value or '')
                        break
            
            if field_data:
                writer.update_page_form_field_values(writer.pages[0], field_data)
        
        pdf_io = io.BytesIO()
        writer.write(pdf_io)
        pdf_io.seek(0)
        return pdf_io.read()
    except Exception as e:
        print(f"Error filling PDF template {pdf_path}: {e}")
        return None

def get_required_hours_for_student(student_user):
    """Get the required internship hours for a student based on their course"""
    try:
        from core.models import CoordinatorSettings
        
        # Get student's course
        if not hasattr(student_user, 'student_profile'):
            return 486  # Default fallback
        
        student_course = student_user.student_profile.course
        
        # Get coordinator settings for the student's college
        # Find a coordinator from the same college
        coordinator = User.objects.filter(
            user_role__role=UserRole.COORDINATOR,
            coordinator_profile__college=student_user.student_profile.college
        ).first()
        
        if not coordinator:
            return 486  # Default fallback
        
        # Get coordinator settings
        settings = CoordinatorSettings.objects.filter(coordinator=coordinator).first()
        
        if not settings or not settings.hours_config:
            return 486  # Default fallback
        
        # Find matching course in hours_config
        # Find matching course in hours_config
        normalized_student_course = student_course.upper().replace('.', '').replace(' ', '')
        
        for config in settings.hours_config:
            program = config.get('program', '').upper()
            normalized_program = program.replace('.', '').replace(' ', '')
            
            # Check for exact match, containment, or common variations
            if (program == student_course.upper() or 
                program in student_course.upper() or 
                student_course.upper() in program or
                normalized_program in normalized_student_course):
                return config.get('requiredHours', 486)
        
        return 486  # Default fallback
    except Exception as e:
        print(f"Error getting required hours: {e}")
        return 486  # Default fallback

def convert_docx_to_pdf(docx_bytes):
    """
    Convert DOCX bytes to PDF bytes.
    On Windows, it tries to use docx2pdf (MS Word) for 1:1 fidelity.
    Falls back to reportlab (text-only) if docx2pdf is unavailable or fails.
    """
    if HAS_DOCX2PDF:
        try:
            # DEBUG LOG
            with open("pdf_debug.log", "a") as log:
                log.write(f"[{timezone.now()}] Attempting docx2pdf conversion...\n")
            
            # Initialize COM for this thread
            pythoncom.CoInitialize()
            
            with tempfile.TemporaryDirectory() as tmp_dir:
                docx_path = os.path.abspath(os.path.join(tmp_dir, "temp.docx"))
                pdf_path = os.path.abspath(os.path.join(tmp_dir, "temp.pdf"))
                
                # Write docx bytes to temp file
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                
                # Convert using docx2pdf (requires MS Word)
                # This gives perfect visual fidelity
                convert(docx_path, pdf_path)
                
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        pdf_data = f.read()
                    print("DEBUG: Successfully converted DOCX to PDF using docx2pdf")
                    with open("pdf_debug.log", "a") as log:
                        log.write(f"[{timezone.now()}] Success using docx2pdf!\n")
                    return pdf_data
                else:
                    with open("pdf_debug.log", "a") as log:
                        log.write(f"[{timezone.now()}] Error: pdf_path does not exist after convert call\n")
                    
        except Exception as e:
            print(f"ERROR: docx2pdf conversion failed: {e}. Falling back to reportlab.")
            with open("pdf_debug.log", "a") as log:
                log.write(f"[{timezone.now()}] ERROR: {str(e)}\n")
        finally:
            try:
                pythoncom.CoUninitialize()
            except:
                pass
    else:
        with open("pdf_debug.log", "a") as log:
            log.write(f"[{timezone.now()}] HAS_DOCX2PDF is False\n")

    # FALLBACK: Traditional reportlab conversion (very basic, text only)
    try:
        # Load the DOCX
        docx_io = io.BytesIO(docx_bytes)
        doc = Document(docx_io)
        
        # Create PDF
        pdf_io = io.BytesIO()
        pdf = SimpleDocTemplate(pdf_io, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading1'],
            fontSize=12,
            textColor='red',
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        # Extract text from DOCX and add to PDF
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine style based on content
                # Only apply red header style to the actual EARIST institution name in header (not in MOA body)
                # Skip red styling if it's part of MOA content (contains FIRST PARTY, SECOND PARTY, etc.)
                is_moa_content = any(keyword in para.text for keyword in ['FIRST PARTY', 'SECOND PARTY', 'WHEREAS', 'MEMORANDUM', '(EARIST)'])
                
                if not is_moa_content and ('EULOGIO "AMANG" RODRIGUEZ INSTITUTE' in para.text or 'Republic of the Philippines' in para.text):
                    p = Paragraph(para.text, header_style)
                elif para.text.isupper() and len(para.text) > 10:
                    p = Paragraph(para.text, title_style)
                else:
                    p = Paragraph(para.text, normal_style)
                story.append(p)
                story.append(Spacer(1, 0.1*inch))
                
                # Add page break AFTER closing statement to keep signatures on page 2
                if 'That the above is done and signed voluntarily' in para.text:
                    story.append(PageBreak())
        
        # Build PDF
        pdf.build(story)
        pdf_io.seek(0)
        return pdf_io.read()
    except Exception as e:
        print(f"Error converting DOCX to PDF: {str(e)}")
        return None


def add_earist_header(doc):
    """Add EARIST official header to document"""
    # Header section
    header_section = doc.sections[0].header
    header_para = header_section.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Republic of the Philippines
    run = header_para.add_run("Republic of the Philippines\n")
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    
    # EARIST full name (bold, red)
    run = header_para.add_run('EULOGIO "AMANG" RODRIGUEZ INSTITUTE OF SCIENCE AND TECHNOLOGY\n')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)  # Red color
    run.font.name = 'Arial'
    
    # Address
    run = header_para.add_run("Nagtahan, Sampaloc, Manila\n")
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    
    # College name (bold, gray)
    run = header_para.add_run("COLLEGE OF COMPUTING STUDIES")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(64, 64, 64)  # Dark gray
    run.font.name = 'Arial'

def add_earist_footer(doc):
    """Add EARIST footer to document"""
    footer_section = doc.sections[0].footer
    footer_para = footer_section.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = footer_para.add_run("CCS@earist.edu.ph")
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

def create_endorsement_letter(student, company, coordinator):
    """Generate endorsement letter in DOCX format"""
    doc = Document()
    add_earist_header(doc)
    add_earist_footer(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ENDORSEMENT LETTER FOR ON-THE-JOB TRAINING")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Arial'
    
    # Add spacing
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(timezone.now().strftime('%B %d, %Y'))
    date_para.style = 'Normal'
    
    doc.add_paragraph()
    
    # Addressee
    addressee = doc.add_paragraph()
    run = addressee.add_run(f"{company.name}\n")
    run.font.bold = True
    addressee.add_run(f"{getattr(company, 'address', '[Company Address]')}")
    
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("Dear Sir/Madam,")
    
    doc.add_paragraph()
    
    # Body
    student_name = f"{student.first_name} {student.last_name}"
    student_id = student.student_profile.student_id if hasattr(student, 'student_profile') else "N/A"
    course = student.student_profile.course if hasattr(student, 'student_profile') else "N/A"
    
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(
        f"This is to formally endorse {student_name}, Student ID: {student_id}, "
        f"currently enrolled in {course} at EARIST - College of Computing Studies, "
        f"for On-the-Job Training (OJT) at your esteemed organization.\n\n"
        f"{student_name} has demonstrated excellent academic performance, strong technical skills, "
        f"and professional attitude throughout their studies. We believe that {student_name} possesses "
        f"the necessary competencies and dedication required to contribute meaningfully to your organization "
        f"during the internship period.\n\n"
        f"We hope that you will consider accepting {student_name} as an OJT trainee at {company.name}. "
        f"This internship opportunity will provide valuable practical experience and help bridge the gap "
        f"between academic learning and professional practice."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Closing
    closing = doc.add_paragraph("Respectfully yours,")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature line
    sig = doc.add_paragraph()
    sig.add_run("_" * 40)
    
    # Coordinator name
    coord_name = doc.add_paragraph()
    run = coord_name.add_run(f"{coordinator.first_name} {coordinator.last_name}")
    run.font.bold = True
    
    # Title
    doc.add_paragraph("OJT Coordinator")
    doc.add_paragraph("College of Computing Studies")
    
    return doc

def create_recommendation_letter(student, company, position, coordinator):
    """Generate recommendation letter in DOCX format"""
    doc = Document()
    add_earist_header(doc)
    add_earist_footer(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LETTER OF RECOMMENDATION")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Date
    doc.add_paragraph(timezone.now().strftime('%B %d, %Y'))
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("To Whom It May Concern,")
    doc.add_paragraph()
    
    # Body
    student_name = f"{student.first_name} {student.last_name}"
    student_id = student.student_profile.student_id if hasattr(student, 'student_profile') else "N/A"
    course = student.student_profile.course if hasattr(student, 'student_profile') else "N/A"
    
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(
        f"I am pleased to recommend {student_name} (Student ID: {student_id}), "
        f"who successfully completed their On-the-Job Training (OJT) at {company.name} "
        f"as a {position}.\n\n"
        f"{student_name}, a {course} student from EARIST - College of Computing Studies, "
        f"demonstrated exceptional skills, professionalism, and dedication throughout the internship period. "
        f"{student_name} consistently showed strong technical abilities, excellent work ethic, and the capacity "
        f"to work effectively both independently and as part of a team.\n\n"
        f"I highly recommend {student_name} for future employment opportunities and am confident that "
        f"they will be a valuable asset to any organization."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.add_run("_" * 40)
    
    coord_name = doc.add_paragraph()
    run = coord_name.add_run(f"{coordinator.first_name} {coordinator.last_name}")
    run.font.bold = True
    
    doc.add_paragraph("OJT Coordinator")
    doc.add_paragraph("EARIST - College of Computing Studies")
    
    return doc

def create_certificate(student, company, position, coordinator):
    """Generate certificate of completion in DOCX format"""
    doc = Document()
    add_earist_header(doc)
    add_earist_footer(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CERTIFICATE OF COMPLETION")
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Certificate text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert_text.add_run("This certifies that\n\n")
    
    # Student name (larger, bold)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(f"{student.first_name} {student.last_name}")
    run.font.size = Pt(18)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Details
    student_id = student.student_profile.student_id if hasattr(student, 'student_profile') else "N/A"
    course = student.student_profile.course if hasattr(student, 'student_profile') else "N/A"
    
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(
        f"Student ID: {student_id}\n"
        f"Course: {course}\n\n"
        f"has successfully completed the required OJT hours at\n"
        f"{company.name}\n"
        f"as a {position}"
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Date Issued: {timezone.now().strftime('%B %d, %Y')}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run("_" * 40)
    
    coord_name = doc.add_paragraph()
    coord_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = coord_name.add_run(f"{coordinator.first_name} {coordinator.last_name}")
    run.font.bold = True
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run("OJT Coordinator, EARIST")
    
    return doc

def create_acceptance_letter(company, coordinator):
    """Generate acceptance letter for company to accept students"""
    doc = Document()
    add_earist_header(doc)
    add_earist_footer(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ACCEPTANCE LETTER FOR OJT PROGRAM")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(timezone.now().strftime('%B %d, %Y'))
    
    doc.add_paragraph()
    
    # Addressee
    addressee = doc.add_paragraph()
    run = addressee.add_run(f"{company.name}\n")
    run.font.bold = True
    addressee.add_run(f"{getattr(company, 'address', '[Company Address]')}")
    
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("Dear Sir/Madam,")
    
    doc.add_paragraph()
    
    # Body
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.add_run(
        f"Greetings from EARIST - College of Computing Studies!\n\n"
        f"We are pleased to inform you that {company.name} has been selected as one of our "
        f"partner companies for our On-the-Job Training (OJT) program. We would like to formally "
        f"request your acceptance of our students for their internship training.\n\n"
        f"Our students are well-prepared and eager to gain practical experience in the industry. "
        f"We believe that this partnership will be mutually beneficial and will provide our students "
        f"with valuable hands-on experience in their field of study.\n\n"
        f"We look forward to a fruitful collaboration with your esteemed organization."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Respectfully yours,")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature line
    sig = doc.add_paragraph()
    sig.add_run("_" * 40)
    
    # Coordinator name
    coord_name = doc.add_paragraph()
    run = coord_name.add_run(f"{coordinator.first_name} {coordinator.last_name}")
    run.font.bold = True
    
    # Title
    doc.add_paragraph("OJT Coordinator")
    doc.add_paragraph("College of Computing Studies")
    
    return doc

def create_waiver_form(student, company, coordinator):
    """Generate Parent/Guardian Consent Form in DOCX format"""
    doc = Document()
    add_earist_header(doc)
    add_earist_footer(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PARENT/GUARDIAN CONSENT FORM")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # AY, Term (right-aligned)
    ay_para = doc.add_paragraph()
    ay_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ay_para.add_run("AY, Term\n______________")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Parent/Guardian statement
    student_name = f"{student.first_name} {student.last_name}"
    
    intro = doc.add_paragraph()
    intro.add_run("I, _________________________________________________, the parent/legal guardian\n")
    intro.add_run("of ")
    run = intro.add_run(student_name)  # Auto-fill student name here
    run.font.bold = True
    intro.add_run(", hereby expressly state that I agree to the following:")
    
    doc.add_paragraph()
    
    # Get student's course
    student_course = "Bachelor in Multimedia Arts"  # Default
    if hasattr(student, 'student_profile') and student.student_profile and student.student_profile.course:
        student_course = student.student_profile.course
    
    company_name = company.name if company and hasattr(company, 'name') else "<name of Host Training Establishment>"
    
    # Point 1
    p1 = doc.add_paragraph()
    p1.add_run("    1. To allow my son/daughter, ")
    run = p1.add_run("<name of Student Trainee>")  # Placeholder
    run.font.bold = True
    p1.add_run(" to take his/her On-the-Job\n")
    p1.add_run("       Training (OJT) at ")
    run = p1.add_run(company_name)  # Auto-fill company name
    run.font.bold = True
    p1.add_run(" for ")
    run = p1.add_run("______")  # Blank line instead of 300
    run.font.bold = True
    p1.add_run(" hours in partial\n")
    p1.add_run("       fulfillment of the requirements for the degree in ")
    run = p1.add_run(student_course)  # Use actual student course
    run.font.bold = True
    p1.add_run(".")
    
    doc.add_paragraph()
    
    # Point 2 (normal black text, EARIST in bold)
    p2 = doc.add_paragraph()
    run = p2.add_run("    2. I have read the rules and regulations set by the ")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p2.add_run("EARIST")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p2.add_run(" OJT Course Policy and the Host\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p2.add_run("       Training Establishment and commits that my son/daughter will abide by the said rules\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p2.add_run("       and regulations.")
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Point 3 (normal black text, EARIST in bold)
    p3 = doc.add_paragraph()
    run = p3.add_run("    3. I fully and voluntarily waive my right to hold ")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p3.add_run("EARIST")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p3.add_run(" ")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p3.add_run("<Name of Campus>")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p3.add_run(" and/or any of its\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p3.add_run("       officers, employees, or representatives responsible for any case of untoward incident\n")
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = p3.add_run("       that may happen to my son/daughter during the duration of his/her training.")
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Parent/Guardian signature line
    sig1_line = doc.add_paragraph()
    sig1_line.add_run("_" * 60 + "          " + "_" * 20)
    
    sig1_label = doc.add_paragraph()
    sig1_label.add_run("Signature over Printed Name of Parent or Guardian Date Signed          Signature")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Received by
    doc.add_paragraph("Received by:")
    
    doc.add_paragraph()
    
    # OJT Adviser signature line
    sig2_line = doc.add_paragraph()
    sig2_line.add_run("_" * 60 + "          " + "_" * 20)
    
    sig2_label = doc.add_paragraph()
    sig2_label.add_run("Signature over Printed Name of OJT Adviser Date Signed          Signature")
    
    return doc

# ========================================
# COORDINATOR DASHBOARD
# ========================================

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_dashboard(request):
    try:
        # Get coordinator's college
        coordinator_college = None
        if hasattr(request.user, 'coordinator_profile'):
            coordinator_college = request.user.coordinator_profile.college
        
        # If admin or no college assigned, show all data
        is_admin = hasattr(request.user, 'user_role') and request.user.user_role.role == 'admin'
        
        if is_admin or not coordinator_college:
            # Admin sees everything
            total_students = StudentProfile.objects.count()
            pending_applications = Application.objects.filter(status='Pending').count()
            
            # Active internships = approved applications
            active_internships = Application.objects.filter(status='Approved').count()
            
            recent_applications = Application.objects.all().order_by('-applied_at')[:5]
            
            # Admin sees all companies
            total_companies = Company.objects.count()
        else:
            # Coordinator sees only their college's data
            # Total students from their college
            total_students = StudentProfile.objects.filter(college=coordinator_college).count()
            
            # Get student user IDs from this college
            college_student_ids = StudentProfile.objects.filter(
                college=coordinator_college
            ).values_list('user_id', flat=True)
            
            # Pending applications from students in this college
            pending_applications = Application.objects.filter(
                student_id__in=college_student_ids,
                status='Pending'
            ).count()
            
            # Active internships = approved applications from students in this college
            active_internships = Application.objects.filter(
                student_id__in=college_student_ids,
                status='Approved'
            ).count()
            
            # Recent applications from this college's students
            recent_applications = Application.objects.filter(
                student_id__in=college_student_ids
            ).order_by('-applied_at')[:5]
            
            # Companies: approved companies targeted to this college OR with students from this college
            # Get unique company IDs from applications of students in this college
            company_ids_with_students = set(Application.objects.filter(
                student_id__in=college_student_ids,
                internship__company__isnull=False
            ).values_list('internship__company', flat=True).distinct())
            
            # Get all approved companies and filter by target_colleges in Python
            approved_companies = Company.objects.filter(status='Approved')
            targeted_company_ids = set()
            for company in approved_companies:
                if company.target_colleges and coordinator_college in company.target_colleges:
                    targeted_company_ids.add(company.id)
            
            # Combine both sets
            all_company_ids = company_ids_with_students | targeted_company_ids
            total_companies = len(all_company_ids)
        
        return Response({
            'total_students': total_students,
            'total_companies': total_companies,
            'pending_applications': pending_applications,
            'active_internships': active_internships,
        })
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# ========================================
# COMPANY MANAGEMENT
# ========================================

@api_view(['PUT'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_approve_company(request, company_id):
    try:
        company = Company.objects.get(id=company_id)
        action = request.data.get('action')
        
        if action == 'approve':
            if hasattr(company, 'is_approved'):
                company.is_approved = True
            company.status = 'Approved'
            company.approved_by = request.user
            company.approved_at = timezone.now()
            company.save()
            return Response({'message': f'Company {company.name} approved', 'company': CompanySerializer(company).data})
        elif action == 'reject':
            if hasattr(company, 'is_approved'):
                company.is_approved = False
            company.status = 'Rejected'
            company.rejection_reason = request.data.get('reason', '')
            company.save()
            return Response({'message': f'Company {company.name} rejected', 'company': CompanySerializer(company).data})
        else:
            return Response({'error': 'Invalid action'}, status=status.HTTP_400_BAD_REQUEST)
    except Company.DoesNotExist:
        return Response({'error': 'Company not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

# ========================================
# DOCUMENT GENERATION
# ========================================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_generate_document(request):
    """Generate professional DOCX documents with EARIST branding"""
    print("DEBUG: Received document generation request")
    print(f"DEBUG: Request Data: {request.data}")
    
    try:
        # Define Placeholder class
        class PlaceholderCompany:
            name = "____________________________________"
            address = "______________________________________________________"
            contact_person = "__________________________"

        # Extract Data
        document_type = request.data.get('document_type')
        student_id = request.data.get('student_id')
        company_id = request.data.get('company_id')
        
        if not document_type:
            print("ERROR: Document type missing in request data")
            return Response({'error': 'Document type is required'}, status=status.HTTP_400_BAD_REQUEST)

        # Fetch Entities
        student = None
        student_profile = None
        company = None
        position = "[Internship Position]"
        
        if student_id:
            try:
                student = User.objects.get(id=student_id)
                if hasattr(student, 'student_profile'):
                    student_profile = student.student_profile
                
                # Try to find internship
                application = Application.objects.filter(student=student, status='Approved').first()
                if application:
                    company = application.internship.company
                    position = application.internship.position
            except User.DoesNotExist:
                pass

        if company_id and not company:
            try:
                company = Company.objects.get(id=company_id)
            except Company.DoesNotExist:
                pass
        
        if not company:
            print("DEBUG: No company found, using PlaceholderCompany")
            company = PlaceholderCompany()

        print(f"DEBUG: Processing document_type: '{document_type}' for student: '{student}'")

        # Check for custom template first
        doc = None
        try:
            # Map legacy request codes to database codes if necessary
            db_code = document_type
            legacy_map = {
                'recommendation': 'recommendation_letter',
                'certificate': 'completion_certificate',
                'waiver': 'waiver_consent'
            }
            if document_type in legacy_map:
                db_code = legacy_map[document_type]
            
            doc_config = DocumentTypeConfig.objects.filter(code=db_code).first()
            if not doc_config and db_code != document_type:
                 # Fallback: try original code if mapped one wasn't found
                 doc_config = DocumentTypeConfig.objects.filter(code=document_type).first()
                 
            if doc_config and doc_config.template_file:
                print(f"DEBUG: Found custom template for {document_type}: {doc_config.template_file.path}")
                
                # Get required hours
                req_hours = get_required_hours_for_student(student) if student else 486
                
                # Prepare Context for replacement with extra fields
                # Format section to match EARIST standard: "4-C BS INFOTECH"
                section_display = ""
                if student_profile:
                    # Extract year number (e.g., "4" from "4th Year")
                    year_num = ""
                    if student_profile.year:
                        year_str = student_profile.year.lower()
                        if '1st' in year_str or 'first' in year_str:
                            year_num = "1"
                        elif '2nd' in year_str or 'second' in year_str:
                            year_num = "2"
                        elif '3rd' in year_str or 'third' in year_str:
                            year_num = "3"
                        elif '4th' in year_str or 'fourth' in year_str:
                            year_num = "4"
                    
                    # Get section letter
                    section_letter = student_profile.section if student_profile.section else ""
                    
                    # Get course abbreviation (e.g., "BS INFOTECH" from full course name)
                    course_abbr = ""
                    if student_profile.course:
                        course = student_profile.course
                        # Extract abbreviation from parentheses if it exists
                        if '(' in course and ')' in course:
                            course_abbr = course[course.find('(')+1:course.find(')')].replace('.', '')
                        else:
                            # Fallback: use full course name
                            course_abbr = course
                    
                    # Combine: "4-C BS INFOTECH"
                    if year_num and section_letter and course_abbr:
                        section_display = f"{year_num}-{section_letter} {course_abbr}"
                    elif section_letter and course_abbr:
                        section_display = f"{section_letter} {course_abbr}"
                    elif year_num and section_letter:
                        section_display = f"{year_num}-{section_letter}"
                    elif section_letter:
                        section_display = section_letter
                
                # Check if student has an approved company
                has_company = company and hasattr(company, 'name') and not company.name.startswith('_')
                
                context = {
                    'student_name': f"{student.first_name} {student.last_name}" if student else "__________________________",
                    'student_first_name': student.first_name if student else "________________",
                    'student_last_name': student.last_name if student else "________________",
                    'student_id': student_profile.student_id if (student_profile and student_profile.student_id) else "________________",
                    'course': student_profile.course if (student_profile and student_profile.course) else "________________",
                    'college': student_profile.get_college_display() if (student_profile and student_profile.college) else "________________",
                    'section': section_display if section_display else "________________",
                    'year': student_profile.year if (student_profile and student_profile.year) else "________",
                    'company_name': company.name if has_company else "______________________________",
                    'company_address': company.address if (has_company and hasattr(company, 'address')) else "______________________________",
                    'company_contact': company.contact_person if (has_company and hasattr(company, 'contact_person')) else "__________________________",
                    'position': position if position != "[Internship Position]" else "________________",
                    'required_hours': str(req_hours),
                    'hours': str(req_hours),
                    'coordinator_name': f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username,
                    'date': timezone.now().strftime('%B %d, %Y'),
                    'current_date': timezone.now().strftime('%B %d, %Y'),
                    'current_year': timezone.now().strftime('%Y'),
                    'academic_year': f"{timezone.now().year}-{timezone.now().year + 1}",
                }
                
                doc = fill_document_template(doc_config.template_file.path, context)
        except Exception as e:
            print(f"ERROR loading custom template: {e}")
            # Fallback if custom template fails
            doc = None

        # Generate Document based on type if no custom template used
        if doc:
            pass
            
        elif document_type == 'endorsement_letter':
            if not student or not student_profile:
                return Response({'error': 'Student required for endorsement'}, status=status.HTTP_400_BAD_REQUEST)
            doc = create_endorsement_letter(student, company, request.user)
            
        elif document_type == 'recommendation':
            if not student or not student_profile:
                return Response({'error': 'Student required for recommendation'}, status=status.HTTP_400_BAD_REQUEST)
            doc = create_recommendation_letter(student, company, position, request.user)
            
        elif document_type == 'certificate':
            if not student:
                return Response({'error': 'Student selection is required for certificate'}, status=status.HTTP_400_BAD_REQUEST)
            if not student_profile:
                return Response({'error': f'Student {student.get_full_name()} does not have a profile.'}, status=status.HTTP_400_BAD_REQUEST)
            doc = create_certificate(student, company, position, request.user)
            
        elif document_type == 'acceptance_letter':
            # Acceptance letter only requires company, not student
            doc = create_acceptance_letter(company, request.user)
            
        elif document_type == 'waiver' or document_type == 'consent_letter':
            if not student:
                return Response({'error': 'Student selection is required for waiver/consent'}, status=status.HTTP_400_BAD_REQUEST)
            if not student_profile:
                return Response({'error': f'Student {student.get_full_name()} does not have a profile. Please ensure they have completed their profile setup.'}, status=status.HTTP_400_BAD_REQUEST)
            doc = create_waiver_form(student, company, request.user)
            
        elif document_type == 'contract_moa':
            # Contract/MOA only requires company
            doc = create_moa_contract(company, request.user)
            
        else:
            print(f"ERROR: Document type '{document_type}' not recognized by backend")
            return Response({
                'error': f'Document type "{document_type}" not yet implemented on the server.',
                'received_type': document_type,
                'available_types': ['endorsement_letter', 'recommendation', 'certificate', 'acceptance_letter', 'waiver', 'consent_letter', 'contract_moa']
            }, status=status.HTTP_400_BAD_REQUEST)

        # Get requested format (default to PDF)
        output_format = request.data.get('format', 'pdf').lower()
        
        # Check for custom PDF template if PDF requested
        pdf_bytes = None
        if output_format == 'pdf' and 'doc_config' in locals() and doc_config and doc_config.template_file_pdf:
            print(f"DEBUG: Using custom PDF template for {document_type}")
            pdf_bytes = fill_pdf_template(doc_config.template_file_pdf.path, context)
            if not pdf_bytes:
                # Fallback: if filling fails, just return the raw PDF if requested specifically by user
                with open(doc_config.template_file_pdf.path, 'rb') as f:
                    pdf_bytes = f.read()

        # Save document to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        docx_bytes = doc_io.read()
        
        # Create DocumentTemplate record
        doc_name = document_type.replace('_', ' ').title()
        if student:
            doc_name += f" - {student.first_name} {student.last_name}"
        
        timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
        
        # Save DOCX version
        safe_student_name = f"{student.first_name}_{student.last_name}".replace(' ', '_') if student else "Template"
        docx_file_name = f"{document_type}_{safe_student_name}_{timestamp}.docx"
        docx_content = ContentFile(docx_bytes)
        
        document = DocumentTemplate.objects.create(
            name=doc_name,
            template_type=document_type,
            uploaded_by=request.user,
            description=f"Generated document for {doc_name}"
        )
        
        document.file.save(docx_file_name, docx_content)
        
        # Generate and save PDF version if not already generated from template
        if not pdf_bytes:
            pdf_bytes = convert_docx_to_pdf(docx_bytes)
            
        if pdf_bytes:
            pdf_file_name = f"{document_type}_{safe_student_name}_{timestamp}.pdf"
            pdf_path = document.file.path.replace('.docx', '.pdf')
            # If we used a custom PDF template, the filename inside the DB folder should match the DOCX
            with open(pdf_path, 'wb') as pdf_file:
                pdf_file.write(pdf_bytes)
        
        document.save()
        
        # Determine which file to return
        if output_format == 'docx':
            file_url = document.file.url if document.file else None
            # Use media-view endpoint for compatible previewing/embedding
            download_url = f'/api/media-view/{document.file.name}' if document.file else None
            file_format = 'docx'
        else:  # Default to PDF
            pdf_path_relative = document.file.name.replace(".docx", ".pdf")
            pdf_url = f'{settings.MEDIA_URL}{pdf_path_relative}'
            # Use media-view endpoint for compatible previewing/embedding
            download_url = f'/api/media-view/{pdf_path_relative}'
            file_url = pdf_url
            file_format = 'pdf'
        
        response_data = {
            'message': f'{doc_name} generated successfully',
            'document_id': document.id,
            'file_url': file_url,
            'download_url': download_url,
            'pdf_url': f'/api/media-view/{document.file.name.replace(".docx", ".pdf")}',
            'docx_url': f'/api/media-view/{document.file.name}',
            'format': file_format,
            'name': doc_name
        }
        
        if student:
            response_data['student_name'] = f"{student.first_name} {student.last_name}"
        if company and not isinstance(company, PlaceholderCompany):
            response_data['company_name'] = company.name
            
        return Response(response_data, status=status.HTTP_201_CREATED)

    except Exception as e:
        print(f"CRITICAL ERROR in document generation: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

# ========================================
# STUDENT MANAGEMENT
# ========================================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_bulk_approve_students(request):
    try:
        student_ids = request.data.get('student_ids', [])
        if not student_ids:
            return Response({'error': 'No student IDs provided'}, status=status.HTTP_400_BAD_REQUEST)
        
        approved_count = 0
        for student_id in student_ids:
            try:
                user = User.objects.get(id=student_id)
                user.is_active = True
                user.save()
                
                # Approve all pre-training requirements
                from .models import PreTrainingRequirement
                requirements = PreTrainingRequirement.objects.filter(student=user)
                for req in requirements:
                    req.status = 'Approved'
                    req.reviewed_at = timezone.now()
                    req.admin_comment = 'Your document has been verified'
                    req.save()
                
                if hasattr(user, 'student_profile'):
                    profile = user.student_profile
                    profile.is_approved = True
                    profile.cor_verified = True  # Also verify COR
                    profile.approved_by = request.user
                    profile.approved_at = timezone.now()
                    profile.save()
                approved_count += 1
            except User.DoesNotExist:
                continue
        return Response({'message': f'Approved {approved_count} students', 'approved_count': approved_count})
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_users_list(request):
    try:
        students = User.objects.filter(user_role__role='student')
        if hasattr(request.user, 'user_role') and request.user.user_role.role == 'coordinator':
            if hasattr(request.user, 'coordinator_profile'):
                college = request.user.coordinator_profile.college
                if college:
                    students = students.filter(student_profile__college=college)
        serializer = UserSerializer(students, many=True)
        return Response(serializer.data)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# ========================================
# SETTINGS
# ========================================

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_student_grades(request):
    """Get all student grades for coordinator's college"""
    try:
        from .models import StudentFinalGrade
        
        # Get coordinator's college
        coordinator_college = None
        if hasattr(request.user, 'coordinator_profile'):
            coordinator_college = request.user.coordinator_profile.college
        
        # Get students from coordinator's college
        students = User.objects.filter(user_role__role='student')
        if hasattr(request.user, 'user_role') and request.user.user_role.role == 'coordinator':
            if coordinator_college:
                students = students.filter(student_profile__college=coordinator_college)
        
        # Get all grades for these students
        student_ids = list(students.values_list('id', flat=True))
        grades = StudentFinalGrade.objects.filter(student_id__in=student_ids)
        
        # Format response
        grades_data = {}
        for grade in grades:
            grades_data[grade.student_id] = {
                'student': grade.student.username,
                'attendance_score': float(grade.attendance_score),
                'supervisor_score': float(grade.supervisor_rating_score),
                'requirements_score': float(grade.requirements_score),
                'final_grade': float(grade.final_grade),
                'remarks': grade.remarks
            }
        
        return Response(grades_data)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET', 'PUT'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_settings(request):
    try:
        if request.method == 'GET':
            coordinator_profile = request.user.coordinator_profile if hasattr(request.user, 'coordinator_profile') else None
            
            # Get or create coordinator settings
            from core.models import CoordinatorSettings
            settings, created = CoordinatorSettings.objects.get_or_create(
                coordinator=request.user,
                defaults={
                    'hours_config': [],
                    'required_docs': [],
                    'cutoff_dates': [],
                    'activities': []
                }
            )
            
            return Response({
                'first_name': request.user.first_name,
                'last_name': request.user.last_name,
                'email': request.user.email,
                'college': coordinator_profile.college if coordinator_profile else '',
                'department': coordinator_profile.department if coordinator_profile else '',
                'phone': coordinator_profile.phone if coordinator_profile else '',
                'office_location': coordinator_profile.office_location if coordinator_profile else '',
                'hours_config': settings.hours_config or [],
                'required_docs': settings.required_docs or [],
                'cutoff_dates': settings.cutoff_dates or [],
                'activities': settings.activities or []
            })
        elif request.method == 'PUT':
            user = request.user
            user.first_name = request.data.get('first_name', user.first_name)
            user.last_name = request.data.get('last_name', user.last_name)
            user.save()
            
            if hasattr(user, 'coordinator_profile'):
                profile = user.coordinator_profile
                profile.college = request.data.get('college', profile.college)
                profile.department = request.data.get('department', profile.department)
                profile.phone = request.data.get('phone', profile.phone)
                profile.office_location = request.data.get('office_location', profile.office_location)
                profile.save()
            
            # Save coordinator settings
            from core.models import CoordinatorSettings
            settings, created = CoordinatorSettings.objects.get_or_create(coordinator=user)
            
            if 'hours_config' in request.data:
                settings.hours_config = request.data['hours_config']
            if 'required_docs' in request.data:
                settings.required_docs = request.data['required_docs']
            if 'cutoff_dates' in request.data:
                settings.cutoff_dates = request.data['cutoff_dates']
            if 'activities' in request.data:
                settings.activities = request.data['activities']
            
            settings.save()
            
            return Response({'message': 'Settings updated successfully'})
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
def create_moa_contract(company, coordinator):
    """Generate Memorandum of Agreement between EARIST and Company"""
    doc = Document()
    add_earist_header(doc)
    add_earist_footer(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MEMORANDUM OF AGREEMENT")
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Opening
    opening = doc.add_paragraph()
    opening.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = opening.add_run("KNOW ALL MEN BY THESE PRESENTS:")
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Introduction
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "This memorandum of agreement is entered into by and between the following:"
    )
    
    doc.add_paragraph()
    
    # First Party - EARIST
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p1.add_run("EULOGIO \"AMANG\" RODRIGUEZ INSTITUTE OF SCIENCE AND TECHNOLOGY (EARIST)")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)  # Explicit black
    p1.add_run(
        ", (An educational institution, duly organized and registered "
        "under the laws and regulations of the government located at No. 5, Manuel L. Quezon St., Hagomy, "
        "Taguig City represented by its President, Atty. Ronald R. Yap, referred to as the "
    )
    run = p1.add_run("FIRST PARTY")
    run.font.bold = True
    p1.add_run(",")
    
    doc.add_paragraph()
    
    # "and"
    and_para = doc.add_paragraph()
    and_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    and_para.add_run("and")
    
    doc.add_paragraph()
    
    # Second Party - Company
    company_name = company.name if company and hasattr(company, 'name') else "[Company Name]"
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p2.add_run(company_name)
    run.font.bold = True
    p2.add_run(
        ", a corporation/company duly registered and organized "
        "under the laws and regulations of the government located at "
    )
    p2.add_run("_" * 50)
    p2.add_run("\nrepresented by its President or HR Manager ")
    p2.add_run("_" * 50)
    p2.add_run(" referred to as the ")
    run = p2.add_run("SECOND PARTY")
    run.font.bold = True
    p2.add_run(".")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # WHEREAS clauses
    whereas1 = doc.add_paragraph()
    whereas1.add_run("    1. WHEREAS, the students of the FIRST PARTY need to comply with the curriculum\n")
    whereas1.add_run("       requirements of On-the-Job Training/Internship subject;")
    
    doc.add_paragraph()
    
    whereas2 = doc.add_paragraph()
    whereas2.add_run("    2. WHEREAS, the SECOND PARTY is willing to accept the students of the FIRST PARTY\n")
    whereas2.add_run("       and would like to comply with the curriculum requirements as needed;")
    
    doc.add_paragraph()
    
    whereas3 = doc.add_paragraph()
    whereas3.add_run("    3. WHEREAS, it is agreed that there will be no employee – employer relationship between\n")
    whereas3.add_run("       the students of the FIRST PARTY and the SECOND PARTY;")
    
    doc.add_paragraph()
    
    whereas4 = doc.add_paragraph()
    whereas4.add_run("    4. WHEREAS, the Parents/Guardians of the students of the FIRST PARTY shall likewise\n")
    whereas4.add_run("       sign the agreement to comply with the terms and conditions of this agreement;")
    
    doc.add_paragraph()
    
    whereas5 = doc.add_paragraph()
    whereas5.add_run("    5. WHEREAS, it is understood that the SECOND PARTY shall not be held liable in case of\n")
    whereas5.add_run("       accident of the students of the FIRST PARTY, during the duration of the Training;")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add page break BEFORE closing statement to keep all signatures on page 2
    doc.add_page_break()
    
    # Closing statement
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing.add_run("That the above is done and signed voluntarily by all parties concerned.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature section - EARIST header (left-aligned)
    college_header = doc.add_paragraph()
    run = college_header.add_run("EULOGIO \"AMANG\" RODRIGUEZ INSTITUTE OF SCIENCE AND TECHNOLOGY (EARIST)")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # By line with tabs for alignment
    by1 = doc.add_paragraph()
    by1.add_run("By\t\t\t\t\t\t\t\tBy")
    
    doc.add_paragraph()
    
    # Prof. Marlon B. Raquel
    prof1 = doc.add_paragraph()
    run = prof1.add_run("Prof. Marlon B. Raquel")
    run.font.bold = True
    
    prof1_sig = doc.add_paragraph()
    prof1_sig.add_run("CBAA Chairperson\t" + "_" * 30 + "\t\t" + "_" * 40)
    
    doc.add_paragraph()
    
    prof1_label = doc.add_paragraph()
    prof1_label.add_run("\t\t\t\t\t\t\tPresident/HR Manager")
    
    doc.add_paragraph()
    
    # Mrs. Ma. Meni Kathleen L. Osorio
    prof2 = doc.add_paragraph()
    run = prof2.add_run("Mrs. Ma. Meni Kathleen L. Osorio")
    run.font.bold = True
    
    prof2_title = doc.add_paragraph()
    prof2_title.add_run("Acting College Director/College Registrar\t\t\tWith my Conformity:")
    
    doc.add_paragraph()
    
    conformity = doc.add_paragraph()
    conformity.add_run("\t\t\t\t\t\t\t" + "_" * 40)
    
    doc.add_paragraph()
    
    # Attested by
    attested = doc.add_paragraph()
    attested.add_run("Attested by:")
    
    doc.add_paragraph()
    
    # Atty. Ronald R. Yap
    board = doc.add_paragraph()
    run = board.add_run("Atty. Ronald R. Yap")
    run.font.bold = True
    
    board_sig = doc.add_paragraph()
    board_sig.add_run("President/Chairman of the Board\t" + "_" * 25 + "\t\t" + "_" * 40)
    
    doc.add_paragraph()
    
    parent = doc.add_paragraph()
    parent.add_run("\t\t\t\t\t\t\tParent/Guardian")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Notary section
    notary = doc.add_paragraph()
    notary.add_run("SUBSCRIBED AND SWORN to before me on this _____ day of __________ 2017, affiant\n")
    notary.add_run("exhibited to me his/her Community Tax Certificate ________________ issued on\n")
    notary.add_run("______________ at _______________.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Notary Public
    notary_pub = doc.add_paragraph()
    notary_pub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = notary_pub.add_run("NOTARY PUBLIC")
    run.font.bold = True
    
    # Doc details
    doc_details = doc.add_paragraph()
    doc_details.add_run("Doc. No. __________\n")
    doc_details.add_run("Page No. __________\n")
    doc_details.add_run("Book No. __________\n")
    doc_details.add_run("Series of __________")
    
    return doc

# ========================================
# MESSAGING
# ========================================

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_messages(request):
    """View and send messages for coordinators"""
    if request.method == 'GET':
        try:
            # Get all messages where user is sender or recipient (excluding soft-deleted)
            sent_messages = Message.objects.filter(sender=request.user, deleted_by_sender=False)
            received_messages = Message.objects.filter(recipient=request.user, deleted_by_recipient=False)
            
            messages_data = []
            
            # Format sent messages
            for msg in sent_messages:
                messages_data.append({
                    'id': msg.id,
                    'type': 'sent',
                    'sender': msg.sender.username,
                    'sender_name': f"{msg.sender.first_name} {msg.sender.last_name}" if msg.sender.first_name else msg.sender.username,
                    'recipient': msg.recipient.username,
                    'recipient_name': f"{msg.recipient.first_name} {msg.recipient.last_name}" if msg.recipient.first_name else msg.recipient.username,
                    'subject': msg.subject,
                    'message': msg.message,
                    'attachment': request.build_absolute_uri(msg.attachment.url) if msg.attachment else None,
                    'created_at': msg.created_at,
                    'is_read': msg.is_read,
                    'delivered_at': msg.delivered_at,
                    'read_at': msg.read_at,
                    'reply_to': msg.reply_to.id if msg.reply_to else None
                })
            
            # Format received messages and auto-mark as delivered
            for msg in received_messages:
                # Auto-mark as delivered when fetched
                if not msg.delivered_at:
                    msg.delivered_at = timezone.now()
                    msg.save(update_fields=['delivered_at'])
                
                messages_data.append({
                    'id': msg.id,
                    'type': 'received',
                    'sender': msg.sender.username,
                    'sender_name': f"{msg.sender.first_name} {msg.sender.last_name}" if msg.sender.first_name else msg.sender.username,
                    'recipient': msg.recipient.username,
                    'recipient_name': f"{msg.recipient.first_name} {msg.recipient.last_name}" if msg.recipient.first_name else msg.recipient.username,
                    'subject': msg.subject,
                    'message': msg.message,
                    'attachment': request.build_absolute_uri(msg.attachment.url) if msg.attachment else None,
                    'created_at': msg.created_at,
                    'is_read': msg.is_read,
                    'delivered_at': msg.delivered_at,
                    'read_at': msg.read_at,
                    'reply_to': msg.reply_to.id if msg.reply_to else None
                })
            
            # Sort by created_at
            messages_data.sort(key=lambda x: x['created_at'], reverse=True)
            
            return Response(messages_data)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    elif request.method == 'POST':
        try:
            recipient_id = request.data.get('recipient_id')
            subject = request.data.get('subject')
            message_text = request.data.get('message')
            attachment = request.FILES.get('attachment')
            
            if not all([recipient_id, subject, message_text]):
                return Response({'error': 'Missing required fields'}, status=status.HTTP_400_BAD_REQUEST)
            
            recipient = User.objects.get(id=recipient_id)
            
            reply_to_id = request.data.get('reply_to_id')
            reply_to = None
            if reply_to_id:
                try:
                    reply_to = Message.objects.get(id=reply_to_id)
                except Message.DoesNotExist:
                    pass

            message = Message.objects.create(
                sender=request.user,
                recipient=recipient,
                subject=subject,
                message=message_text,
                attachment=attachment,
                reply_to=reply_to
            )
            
            return Response({
                'id': message.id,
                'created_at': message.created_at,
                'type': 'sent',
                'sender': request.user.username,
                'sender_name': f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username,
                'recipient': recipient.username,
                'subject': message.subject,
                'message': message.message,
                'attachment': request.build_absolute_uri(message.attachment.url) if message.attachment else None,
                'is_read': False,
                'reply_to': message.reply_to.id if message.reply_to else None
            }, status=status.HTTP_201_CREATED)
        except User.DoesNotExist:
            return Response({'error': 'Recipient not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_journals(request):
    """Get journals of students in coordinator's college"""
    try:
        # Get coordinator's college
        coordinator_college = None
        if hasattr(request.user, 'coordinator_profile'):
            coordinator_college = request.user.coordinator_profile.college
        
        # Get students from coordinator's college
        students = User.objects.filter(user_role__role='student')
        if hasattr(request.user, 'user_role') and request.user.user_role.role == 'coordinator':
            if coordinator_college:
                students = students.filter(student_profile__college=coordinator_college)
        
        student_ids = students.values_list('id', flat=True)
        
        # Get journals
        journals = DailyJournal.objects.filter(student_id__in=student_ids)
        
        serializer = DailyJournalSerializer(journals, many=True)
        return Response(serializer.data)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_mark_message_read(request, message_id):
    """Mark a message as read or delete it for coordinators"""
    try:
        if request.method == 'PUT':
            message = Message.objects.get(id=message_id, recipient=request.user)
            message.is_read = True
            message.read_at = timezone.now()
            message.save()
            
            return Response({'message': 'Message marked as read'})
        
        elif request.method == 'DELETE':
            # Allow soft deletion if user is sender or recipient
            message = Message.objects.filter(id=message_id).filter(
                Q(sender=request.user) | Q(recipient=request.user)
            ).first()
            
            if not message:
                return Response({'error': 'Message not found'}, status=status.HTTP_404_NOT_FOUND)
            
            # Soft delete
            if message.sender == request.user:
                message.deleted_by_sender = True
            if message.recipient == request.user:
                message.deleted_by_recipient = True
            message.save()
            
            # If both deleted, remove from DB
            if message.deleted_by_sender and message.deleted_by_recipient:
                message.delete()
                
            return Response({'message': 'Message deleted successfully'})
            
    except Message.DoesNotExist:
        return Response({'error': 'Message not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_delete_all_messages(request):
    """Delete all messages for the current coordinator"""
    try:
        # Soft delete sent messages
        Message.objects.filter(sender=request.user).update(deleted_by_sender=True)
        # Soft delete received messages
        Message.objects.filter(recipient=request.user).update(deleted_by_recipient=True)
        
        # Determine count (approximate) - Just say "Messages deleted"
        deleted_count = Message.objects.filter(
            Q(sender=request.user, deleted_by_sender=True) | 
            Q(recipient=request.user, deleted_by_recipient=True)
        ).count()
        
        # Clean up fully deleted messages
        Message.objects.filter(deleted_by_sender=True, deleted_by_recipient=True).delete()
        
        return Response({
            'message': f'Successfully deleted {deleted_count} messages',
            'count': deleted_count
        })
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


# ========================================
# DOCUMENT TYPE MANAGEMENT
# ========================================

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_document_types(request):
    """List or create document types"""
    if request.method == 'GET':
        try:
            # Get coordinator's college
            coordinator_college = None
            if hasattr(request.user, 'coordinator_profile'):
                coordinator_college = request.user.coordinator_profile.college
            
            # Filter by college or show all if admin/no college
            is_admin = hasattr(request.user, 'user_role') and request.user.user_role.role == 'admin'
            
            if is_admin or not coordinator_college:
                # Admin sees all document types
                document_types = DocumentTypeConfig.objects.all()
            else:
                # Coordinator sees only their college's types or global types (college=null)
                document_types = DocumentTypeConfig.objects.filter(
                    Q(college=coordinator_college) | Q(college__isnull=True)
                )
            
            # Serialize the data
            # List of document types that have code-based templates
            CODE_BASED_TEMPLATES = [
                'endorsement_letter',
                'recommendation_letter', 
                'completion_certificate',
                'acceptance_letter',
                'waiver_consent',
                'consent_letter',
                'contract_moa'
            ]
            
            data = []
            for doc_type in document_types:
                has_uploaded_template = bool(doc_type.template_file)
                has_code_template = doc_type.code in CODE_BASED_TEMPLATES
                
                data.append({
                    'id': doc_type.id,
                    'name': doc_type.name,
                    'code': doc_type.code,
                    'description': doc_type.description,
                    'category': doc_type.category,
                    'requires_student': doc_type.requires_student,
                    'requires_company': doc_type.requires_company,
                    'is_enabled': doc_type.is_enabled,
                    'college': doc_type.college,
                    'created_at': doc_type.created_at,
                    'updated_at': doc_type.updated_at,
                    'template_file': doc_type.template_file.url if doc_type.template_file else None,
                    'template_file_pdf': doc_type.template_file_pdf.url if doc_type.template_file_pdf else None,
                    'has_template': has_uploaded_template or has_code_template or bool(doc_type.template_file_pdf),
                    'has_uploaded_template': has_uploaded_template,
                    'has_uploaded_pdf_template': bool(doc_type.template_file_pdf),
                    'has_code_template': has_code_template,
                })
            
            return Response(data)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    elif request.method == 'POST':
        try:
            # Create new document type
            data = request.data
            
            # Check if code already exists
            if DocumentTypeConfig.objects.filter(code=data.get('code')).exists():
                return Response({'error': 'Document type with this code already exists'}, status=status.HTTP_400_BAD_REQUEST)
            
            # Get coordinator's college for new document types
            coordinator_college = None
            if hasattr(request.user, 'coordinator_profile'):
                coordinator_college = request.user.coordinator_profile.college
            
            doc_type = DocumentTypeConfig.objects.create(
                name=data.get('name'),
                code=data.get('code'),
                description=data.get('description', ''),
                category=data.get('category', 'Letters'),
                requires_student=data.get('requires_student', False),
                requires_company=data.get('requires_company', False),
                is_enabled=data.get('is_enabled', True),
                college=data.get('college') or coordinator_college,
                created_by=request.user
            )
            
            return Response({
                'message': 'Document type created successfully',
                'document_type': {
                    'id': doc_type.id,
                    'name': doc_type.name,
                    'code': doc_type.code,
                    'description': doc_type.description,
                    'category': doc_type.category,
                    'requires_student': doc_type.requires_student,
                    'requires_company': doc_type.requires_company,
                    'is_enabled': doc_type.is_enabled,
                    'college': doc_type.college,
                }
            }, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['PUT', 'DELETE', 'PATCH'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_document_type_detail(request, doc_type_id):
    """Update, delete, or toggle document type"""
    try:
        doc_type = DocumentTypeConfig.objects.get(id=doc_type_id)
        
        if request.method == 'PUT':
            # Update document type
            data = request.data
            
            # Check if code is being changed and if it already exists
            if data.get('code') and data.get('code') != doc_type.code:
                if DocumentTypeConfig.objects.filter(code=data.get('code')).exists():
                    return Response({'error': 'Document type with this code already exists'}, status=status.HTTP_400_BAD_REQUEST)
            
            doc_type.name = data.get('name', doc_type.name)
            doc_type.code = data.get('code', doc_type.code)
            doc_type.description = data.get('description', doc_type.description)
            doc_type.category = data.get('category', doc_type.category)
            doc_type.requires_student = data.get('requires_student', doc_type.requires_student)
            doc_type.requires_company = data.get('requires_company', doc_type.requires_company)
            doc_type.is_enabled = data.get('is_enabled', doc_type.is_enabled)
            doc_type.college = data.get('college', doc_type.college)
            doc_type.save()
            
            return Response({
                'message': 'Document type updated successfully',
                'document_type': {
                    'id': doc_type.id,
                    'name': doc_type.name,
                    'code': doc_type.code,
                    'description': doc_type.description,
                    'category': doc_type.category,
                    'requires_student': doc_type.requires_student,
                    'requires_company': doc_type.requires_company,
                    'is_enabled': doc_type.is_enabled,
                    'college': doc_type.college,
                }
            })
        
        elif request.method == 'DELETE':
            # Delete document type
            doc_type_name = doc_type.name
            doc_type.delete()
            return Response({'message': f'Document type "{doc_type_name}" deleted successfully'})
        
        elif request.method == 'PATCH':
            # Toggle enabled status
            doc_type.is_enabled = not doc_type.is_enabled
            doc_type.save()
            
            return Response({
                'message': f'Document type {"enabled" if doc_type.is_enabled else "disabled"}',
                'document_type': {
                    'id': doc_type.id,
                    'name': doc_type.name,
                    'is_enabled': doc_type.is_enabled,
                }
            })
    
    except DocumentTypeConfig.DoesNotExist:
        return Response({'error': 'Document type not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
@parser_classes([MultiPartParser, FormParser])
def coordinator_upload_document_template(request, doc_type_id):
    """Upload a template file for a document type"""
    try:
        print(f"DEBUG: Uploading template for doc_type_id={doc_type_id}")
        print(f"DEBUG: request.FILES keys: {list(request.FILES.keys())}")
        print(f"DEBUG: request.data keys: {list(request.data.keys())}")
        
        doc_type = DocumentTypeConfig.objects.get(id=doc_type_id)
        
        if 'template_file' not in request.FILES:
            print("DEBUG: template_file not found in request.FILES")
            return Response({'error': 'DEBUG: No template_file key in request.FILES'}, status=status.HTTP_400_BAD_REQUEST)
        
        template_file = request.FILES['template_file']
        file_name = template_file.name.lower()
        
        # Determine if it's a DOCX or PDF
        if file_name.endswith(('.docx', '.doc')):
            # Delete old template if exists
            if doc_type.template_file:
                doc_type.template_file.delete(save=False)
            doc_type.template_file.save(template_file.name, template_file)
            file_type = 'DOCX'
        elif file_name.endswith('.pdf'):
            # Delete old PDF template if exists
            if doc_type.template_file_pdf:
                doc_type.template_file_pdf.delete(save=False)
            doc_type.template_file_pdf.save(template_file.name, template_file)
            file_type = 'PDF'
        else:
            return Response({'error': 'Invalid file extension. Only DOCX or PDF allowed.'}, status=status.HTTP_400_BAD_REQUEST)
            
        doc_type.save()
        
        return Response({
            'message': f'{file_type} Template uploaded successfully',
            'template_url': doc_type.template_file.url if doc_type.template_file else None,
            'template_file_pdf_url': doc_type.template_file_pdf.url if doc_type.template_file_pdf else None
        })
    
    except DocumentTypeConfig.DoesNotExist:
        return Response({'error': 'Document type not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({'error': f'EXCEPTION: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_preview_code_template(request, doc_type_id):
    """Generate a sample document to preview the code-based template"""
    try:
        doc_type = DocumentTypeConfig.objects.get(id=doc_type_id)
        
        # Create placeholder data for preview
        class PlaceholderStudent:
            first_name = "[Student First Name]"
            last_name = "[Student Last Name]"
            class student_profile:
                student_id = "[Student ID]"
                course = "[Course Name]"
        
        class PlaceholderCompany:
            name = "[Company Name]"
            address = "[Company Address]"
        
        # Generate document based on type
        doc = None
        if doc_type.code == 'endorsement_letter':
            doc = create_endorsement_letter(PlaceholderStudent(), PlaceholderCompany(), request.user)
        elif doc_type.code == 'recommendation_letter':
            doc = create_recommendation_letter(PlaceholderStudent(), PlaceholderCompany(), "[Position]", request.user)
        elif doc_type.code == 'completion_certificate':
            doc = create_certificate(PlaceholderStudent(), PlaceholderCompany(), "[Position]", request.user)
        elif doc_type.code == 'acceptance_letter':
            doc = create_acceptance_letter(PlaceholderCompany(), request.user)
        elif doc_type.code == 'waiver_consent':
            doc = create_waiver_form(PlaceholderStudent(), PlaceholderCompany(), request.user)
        elif doc_type.code == 'contract_moa':
            doc = create_moa_contract(PlaceholderCompany(), request.user)
        else:
            return Response({'error': 'No code-based template available for this document type'}, 
                          status=status.HTTP_404_NOT_FOUND)
        
        # Save to BytesIO and return
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        from django.http import HttpResponse
        response = HttpResponse(
            doc_io.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{doc_type.code}_preview.docx"'
        return response
    
    except DocumentTypeConfig.DoesNotExist:
        return Response({'error': 'Document type not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@role_required([UserRole.COORDINATOR, UserRole.ADMIN])
def coordinator_preview_template_pdf(request, doc_type_id):
    """Convert template to PDF for preview"""
    try:
        import tempfile
        import os
        from docx2pdf import convert
        from django.http import FileResponse
        
        doc_type = DocumentTypeConfig.objects.get(id=doc_type_id)
        
        # Generate or get the DOCX file
        if doc_type.template_file:
            # Use uploaded template
            docx_path = doc_type.template_file.path
        else:
            # Generate code-based template
            class PlaceholderStudent:
                first_name = "[Student First Name]"
                last_name = "[Student Last Name]"
                class student_profile:
                    student_id = "[Student ID]"
                    course = "[Course Name]"
            
            class PlaceholderCompany:
                name = "[Company Name]"
                address = "[Company Address]"
            
            # Generate document based on type
            doc = None
            if doc_type.code == 'endorsement_letter':
                doc = create_endorsement_letter(PlaceholderStudent(), PlaceholderCompany(), request.user)
            elif doc_type.code == 'recommendation_letter':
                doc = create_recommendation_letter(PlaceholderStudent(), PlaceholderCompany(), "[Position]", request.user)
            elif doc_type.code == 'completion_certificate':
                doc = create_certificate(PlaceholderStudent(), PlaceholderCompany(), "[Position]", request.user)
            elif doc_type.code == 'acceptance_letter':
                doc = create_acceptance_letter(PlaceholderCompany(), request.user)
            elif doc_type.code == 'waiver_consent':
                doc = create_waiver_form(PlaceholderStudent(), PlaceholderCompany(), request.user)
            elif doc_type.code == 'contract_moa':
                doc = create_moa_contract(PlaceholderCompany(), request.user)
            else:
                return Response({'error': 'No template available'}, status=status.HTTP_404_NOT_FOUND)
            
            # Save to temporary file
            temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_docx.name)
            temp_docx.close()
            docx_path = temp_docx.name
        
        # Convert DOCX to PDF
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf.close()
        
        try:
            convert(docx_path, temp_pdf.name)
            
            # Return PDF file
            response = FileResponse(
                open(temp_pdf.name, 'rb'),
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'inline; filename="{doc_type.code}_preview.pdf"'
            
            # Clean up temporary DOCX if it was generated
            if not doc_type.template_file and os.path.exists(docx_path):
                os.unlink(docx_path)
            
            return response
        except Exception as convert_error:
            # Clean up on error
            if not doc_type.template_file and os.path.exists(docx_path):
                os.unlink(docx_path)
            if os.path.exists(temp_pdf.name):
                os.unlink(temp_pdf.name)
            raise convert_error
    
    except DocumentTypeConfig.DoesNotExist:
        return Response({'error': 'Document type not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'error': f'PDF conversion failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

