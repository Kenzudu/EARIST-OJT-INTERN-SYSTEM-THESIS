def fill_document_template(template_path, context):
    """
    Fill a DOCX template with context variables.
    Replaces {{variable}} placeholders in paragraphs and tables.
    """
    from docx import Document
    doc = Document(template_path)
    
    def replace_text(paragraph, context):
        if not paragraph.text:
            return
            
        # Check for each key in the context
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                # Simple replacement for now - this might break formatting if runs are split
                # A more robust solution would be needed for complex formatting preservation
                # But for simple templates, this often works
                
                # Check directly in the full text first (simplest case)
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Iterate through all paragraphs in the body
    for paragraph in doc.paragraphs:
        replace_text(paragraph, context)
            
    # Iterate through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text(paragraph, context)
                    
    return doc
